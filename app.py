import streamlit as st
import google.generativeai as genai
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import io

# Set page config
st.set_page_config(
    page_title="Altera JHA Assistant",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="collapsed"  # Changed to collapsed to reduce the sidebar
)

# Configure API
API_KEY = "AIzaSyCDzoVnI44kLyOIDNM4-2h7VzSdtK5gF8g"  # Your new API key
genai.configure(api_key=API_KEY)

# Try to get an available model
try:
    model = genai.GenerativeModel('gemini-1.5-pro')
except Exception as e1:
    try:
        model = genai.GenerativeModel('gemini-1.0-pro')
    except Exception as e2:
        try:
            model = genai.GenerativeModel('gemini-pro')
        except Exception as e3:
            st.error(f"Failed to initialize Gemini model. Please check your API key and connection.")
            st.stop()

# Initialize session state
if 'jha_data' not in st.session_state:
    st.session_state.jha_data = None
if 'equipment_manuals' not in st.session_state:
    st.session_state.equipment_manuals = []
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Define maintenance interval patterns
maintenance_patterns = {
    "daily": r"(?:daily|each day|24\s*hour|every\s*day)",
    "weekly": r"(?:weekly|each week|7\s*day|every\s*week)",
    "monthly": r"(?:monthly|each month|30\s*day|every\s*month)",
    "quarterly": r"(?:quarterly|every\s*3\s*months?|three\s*monthly)",
    "semi_annual": r"(?:semi[- ]annual|every\s*6\s*months?|six\s*monthly)",
    "annual": r"(?:annual|yearly|each year|every\s*year|12\s*month)",
    "two_yearly": r"(?:two[- ]yearly|every\s*2\s*years?|24\s*month)",
    "five_yearly": r"(?:five[- ]yearly|every\s*5\s*years?|60\s*month)"
}

# Work type categories
work_categories = {
    "Working at Height": ["Working Aloft", "Overside"],
    "Critical Equipment": ["Critical Equipment", "Critical System"],
    "Enclosed Space": ["Enclosed Space Entry"],
    "Hot/Cold Work": ["Hot Work", "Cold Work"],
    "Isolation": ["LOTO"]
}

# Functions
def get_work_type_requirements(work_types):
    """Get requirements based on selected work types - simplified version"""
    requirements = {
        "procedures": [],
        "permits": [],
        "special_considerations": [],
        "equipment": [],
        "roles": []
    }
    
    for work_type in work_types:
        if "LOTO" in work_type:
            requirements["procedures"].append("LOTO Procedure: Lock out, tag out all energy sources")
            requirements["roles"].append("Isolation Officer")
            requirements["special_considerations"].append("Verify zero energy state before work")
            
        elif "Working Aloft" in work_type or "Overside" in work_type:
            requirements["equipment"].append("Fall arrest harness and equipment")
            requirements["procedures"].append("Working at Height Procedure: Maintain 100% tie-off")
            
        elif "Critical" in work_type:
            requirements["procedures"].append("Critical Equipment Procedure")
            requirements["roles"].append("Technical Responsible Person")
            
        elif "Enclosed Space" in work_type:
            requirements["equipment"].append("Gas detector and ventilation equipment")
            requirements["procedures"].append("Entry/Exit Control Procedure")
            
        elif any(hw in work_type for hw in ["Hot Work", "Welding", "Cutting", "Grinding"]):
            requirements["procedures"].append("Hot Work Procedure and Fire Watch")
            requirements["special_considerations"].append("Fire sensor inhibition may be required")
            
    return requirements

def get_location_requirements(location):
    """Get requirements specific to the location - simplified version"""
    location_info = {
        "access_requirements": {},
        "hazards": [],
        "controls": [],
        "special_considerations": []
    }
    
    # Define some basic location hazards
    if location == "Engine Room":
        location_info["hazards"] = ["Noise exposure", "Hot surfaces", "Moving machinery"]
        location_info["controls"] = ["Hearing protection", "Appropriate PPE", "Maintain safe distance from machinery"]
    elif location == "Tank Entry":
        location_info["hazards"] = ["Oxygen deficiency", "Toxic atmosphere", "Confined space"]
        location_info["controls"] = ["Gas testing before entry", "Ventilation", "Confined space trained personnel"]
        location_info["access_requirements"] = {"permit": "Required", "gas_test": "Required"}
    elif location in ["MOB", "FFLB", "Gangway"]:
        location_info["hazards"] = ["Working over water", "Fall hazard", "Weather conditions"]
        location_info["controls"] = ["Life jackets", "Fall arrest equipment", "Weather assessment"]
    
    return location_info

def get_risk_badge(consequence, likelihood):
    """Create a colored risk badge based on the risk level"""
    risk_value = consequence * likelihood
    
    if risk_value <= 4:
        return f'<span class="risk-level risk-low">C={consequence} × L={likelihood} = {risk_value}</span>'
    elif risk_value <= 9:
        return f'<span class="risk-level risk-medium">C={consequence} × L={likelihood} = {risk_value}</span>'
    else:
        return f'<span class="risk-level risk-high">C={consequence} × L={likelihood} = {risk_value}</span>'

def create_jha(job_description, work_types, location=None):
    """Generate JHA based on job description, work types, and location"""
    try:
        # Get requirements
        requirements = get_work_type_requirements(work_types)
        
        # Get location-specific requirements
        location_info = get_location_requirements(location) if location else {}
        
        # Build enhanced prompt
        prompt = f"""Create a detailed Job Hazard Analysis (JHA) in JSON format for:

Task: {job_description}

MANDATORY:
Every job MUST start with these three mandatory steps in this order:
1. Job Planning - Daily Work Planning Meeting (reference: RTHB Page 13)
2. Take-5 Assessment (reference: RTHB Page 16)
3. Toolbox Talk (reference: RTHB Page 16)

IMPORTANT REFERENCE GUIDELINES:
- NEVER reference internal code files like planning.py, aloft.py, etc.
- All PPE references should cite "PPE Matrix REF-1412"
- All planning activities should reference "RTHB Page 13"
- All Toolbox Talk and Take-5 activities should reference "RTHB Page 16"

Selected Work Types:
{', '.join(work_types)}

Location: {location if location else 'Not specified'}

Required Procedures:
{chr(10).join(str(proc) for proc in requirements['procedures'])}

Required Roles:
{chr(10).join(str(role) for role in requirements['roles'])}

Equipment Requirements:
{chr(10).join(str(eq) for eq in requirements['equipment'])}

Location-Specific Requirements:
Access Requirements: {json.dumps(location_info.get('access_requirements', {}), indent=2)}
Location Hazards: {chr(10).join(location_info.get('hazards', []))}
Location Controls: {chr(10).join(location_info.get('controls', []))}
Location Considerations: {chr(10).join(location_info.get('special_considerations', []))}

Special Considerations:
{chr(10).join(str(consid) for consid in requirements['special_considerations'])}

Create a comprehensive step-by-step breakdown of the task. For each step:
1. Provide detailed instructions on how to perform the task safely
2. Break down complex actions into separate steps (not sub-steps)
3. Include all necessary preparations and verifications
4. Specify exact tools and equipment needed
5. List all required PPE referencing "PPE Matrix REF-1412"
6. Include communication requirements
7. Add checkpoints and verifications

You MUST include all applicable location-specific requirements in the steps where they apply.
For this {location} location:
1. Include the access requirements in the first step
2. Add location-specific hazards to relevant steps
3. Include location-specific controls in all applicable steps
4. Add location special considerations to the special considerations section

Return ONLY a JSON object with the required structure:
{{
    "steps": [
        {{
            "description": "Detailed step-by-step instructions including:\\n- Required preparations\\n- Tools needed\\n- Specific PPE required\\n- Communication requirements\\n- Verification points",
            "hazards": {{
                "potential_hazards": "List ALL potential hazards for this step",
                "who_affected": "Who or what could be harmed",
                "how_occurs": "Detailed explanation of how harm could occur"
            }},
            "controls": [
                "Specific control measure 1 with correct RTHB page reference",
                "Specific control measure 2 with exact requirements",
                "Required PPE with reference to PPE Matrix REF-1412",
                "Communication protocols",
                "Emergency response measures"
            ],
            "risk_level": {{
                "consequence": 2,
                "likelihood": 2
            }}
        }}
    ],
    "permits_required": ["Required permit 1", "Required permit 2"],
    "special_considerations": {{
        "location_specific": "Location-specific considerations",
        "work_type_specific": "Work type considerations"
    }}
}}

For risk levels, use:
- consequence: 1 (minor) to 5 (catastrophic)
- likelihood: 1 (rare) to 5 (almost certain)"""

        # Generate response
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Remove code block markers if present
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
            
        # Handle potential leading/trailing text
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}')
        if start_idx != -1 and end_idx != -1:
            response_text = response_text[start_idx:end_idx+1]
            
        # Parse the JSON
        parsed_jha = json.loads(response_text)
        
        # Add metadata to JHA
        if location:
            parsed_jha['location'] = location
            
        return parsed_jha
        
    except Exception as e:
        st.error(f"Error generating JHA: {str(e)}")
        # Return a basic error structure
        return {
            "steps": [{
                "description": "Error generating JHA",
                "hazards": {
                    "potential_hazards": "Error in generating hazards",
                    "who_affected": "System error",
                    "how_occurs": "Please try again"
                },
                "controls": ["Please try regenerating the JHA"],
                "risk_level": {"consequence": 1, "likelihood": 1}
            }],
            "permits_required": [],
            "special_considerations": {}
        }

def update_jha_with_message(current_jha, message):
    """Update JHA based on user message"""
    try:
        # Create revision prompt
        revision_prompt = f"""Please revise the existing JHA based on this request: {message}

        Current JHA data:
        {json.dumps(current_jha, indent=2)}

        Important:
        1. Keep all existing steps unless specifically asked to remove them
        2. Maintain the existing structure and format
        3. Only modify what was requested
        4. If there's an error, return the original JHA unchanged
        5. Ensure all hot work activities include fire risk in hazards
    
        Return the complete revised JHA in the same JSON format."""
    
        # Generate revised JHA
        response = model.generate_content(revision_prompt)
        response_text = response.text.strip()
        
        # Clean up response
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
            
        # Handle potential leading/trailing text
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}')
        if start_idx != -1 and end_idx != -1:
            response_text = response_text[start_idx:end_idx+1]
        
        # Parse the revised JHA
        revised_jha = json.loads(response_text)
        
        return revised_jha
        
    except Exception as e:
        st.error(f"Error updating JHA: {str(e)}")
        # Return the original JHA if there was an error
        return current_jha

def create_jha_document(jha_data, vessel_name, task_name):
    """Create a JHA document and return it as bytes"""
    try:
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('JOB HAZARD ANALYSIS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add header table
        header_table = doc.add_table(rows=4, cols=2)
        header_table.style = 'Table Grid'
        
        # Fill header data
        headers = [
            ("Vessel name:", vessel_name),
            ("Task name:", task_name),
            ("Date:", datetime.now().strftime('%d.%m.%Y')),
            ("Supervisor:", "Chief Officer / OOW")
        ]
        
        for i, (label, value) in enumerate(headers):
            row = header_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        # Add spacing
        doc.add_paragraph()
        
        # Create main JHA table
        main_table = doc.add_table(rows=1, cols=5)
        main_table.style = 'Table Grid'
        
        # Set header row
        header_texts = [
            "Step No.",
            "Description of task step",
            "Identify the following:\ni) Hazards\nii) Who or what may be injured\niii) How would this happen",
            "Measures that are required to be put in place",
            "Risk\nC×L=R"
        ]
        
        for i, text in enumerate(header_texts):
            cell = main_table.rows[0].cells[i]
            cell.text = text
        
        # Add steps
        for i, step in enumerate(jha_data.get('steps', []), 1):
            row_cells = main_table.add_row().cells
            
            # Step number
            row_cells[0].text = str(i)
            
            # Description
            row_cells[1].text = step.get('description', '')
            
            # Hazards
            hazards = step.get('hazards', '')
            if isinstance(hazards, dict):
                hazards_text = (
                    f"i. {hazards.get('potential_hazards', '')}\n"
                    f"ii. {hazards.get('who_affected', '')}\n"
                    f"iii. {hazards.get('how_occurs', '')}"
                )
                row_cells[2].text = hazards_text
            else:
                row_cells[2].text = str(hazards)
            
            # Controls
            controls = step.get('controls', [])
            controls_text = '\n'.join(f"• {control}" for control in controls) if isinstance(controls, list) else str(controls)
            row_cells[3].text = controls_text
            
            # Risk
            risk = step.get('risk_level', {})
            c = risk.get('consequence', 1)
            l = risk.get('likelihood', 1)
            risk_text = f"C: {c}\nL: {l}\nR: {c * l}"
            row_cells[4].text = risk_text
        
        # Add spacing
        doc.add_paragraph()
        
        # Add approval section
        approval_table = doc.add_table(rows=3, cols=4)
        approval_table.style = 'Table Grid'
        
        # Add approval headers directly
        approval_table.rows[0].cells[0].text = "Master:"
        approval_table.rows[0].cells[2].text = "Chief Officer:"
        
        approval_table.rows[1].cells[0].text = "Name"
        approval_table.rows[1].cells[1].text = "Signature"
        approval_table.rows[1].cells[2].text = "Name"
        approval_table.rows[1].cells[3].text = "Signature"
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    except Exception as e:
        st.error(f"Error creating document: {str(e)}")
        return None

# Custom CSS - Complete version
st.markdown("""
<style>
    /* Base styles */
    .main {
        background-color: #111827;
        padding: 1rem 2rem;
    }
    .stApp {
        background-color: #111827;
    }
    
    /* Remove sidebar */
    [data-testid="stSidebar"] {
        background-color: #111827;
        width: 0px !important;
        min-width: 0px !important;
        flex: 0 !important;
        -webkit-box-flex: 0 !important;
    }
    
    /* Text and typography */
    h1, h2, h3, h4, h5, h6, .stMarkdown, p, label {
        color: white !important;
    }
    h1 {
        font-size: 2.5rem !important;
        font-weight: 700 !important;
        margin-bottom: 1rem !important;
    }
    h3 {
        font-size: 1.5rem !important;
        font-weight: 600 !important;
        margin-top: 1.5rem !important;
        margin-bottom: 1rem !important;
    }
    
    /* Input controls */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select,
    .stMultiSelect > div > div > select {
        background-color: #1F2937;
        color: white;
        border: 1px solid #374151;
        border-radius: 6px;
        padding: 0.75rem;
        font-size: 1rem;
    }
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #3B82F6;
        box-shadow: 0 0 0 1px #3B82F6;
    }
    
    /* Checkbox styling */
    .stCheckbox label {
        color: white !important;
        font-size: 1rem;
    }
    
    /* Button styling */
    .stButton > button {
        background-color: #2563EB;
        color: white;
        font-weight: 600;
        padding: 0.5rem 1.5rem;
        border-radius: 6px;
        border: none;
        transition: all 0.2s ease;
    }
    .stButton > button:hover {
        background-color: #1D4ED8;
        transform: translateY(-1px);
    }
    
    /* Container styling to replace cards */
    [data-testid="stVerticalBlock"] {
        background-color: #1F2937;
        border-radius: 10px;
        padding: 1rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    /* Section headers */
    .stMarkdown h2, .stMarkdown h3 {
        border-bottom: 1px solid #374151;
        padding-bottom: 0.75rem;
        margin-bottom: 1rem;
    }
    
    /* Horizontal separator */
    hr {
        border-color: #374151;
        margin: 1rem 0;
    }
    
    /* Chat styling */
    .chat-message {
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin-bottom: 0.75rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.12);
    }
    .user-message {
        background-color: #3B82F6;
        color: white;
        margin-left: 20%;
        position: relative;
    }
    .assistant-message {
        background-color: #374151;
        color: white;
        margin-right: 20%;
        position: relative;
    }
    
    /* Expander styling for steps */
    .st-expander {
        background-color: #1F2937;
        border-radius: 8px;
        margin-bottom: 0.75rem;
        border: 1px solid #374151;
    }
    .st-expander:hover {
        border-color: #3B82F6;
    }
    .st-expander-content {
        background-color: #111827;
        padding: 1rem;
        border-radius: 0 0 8px 8px;
    }
    
    /* Risk level styling */
    .risk-level {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-weight: 600;
        margin-top: 0.5rem;
    }
    .risk-low {
        background-color: rgba(16, 185, 129, 0.2);
        color: #10B981;
    }
    .risk-medium {
        background-color: rgba(245, 158, 11, 0.2);
        color: #F59E0B;
    }
    .risk-high {
        background-color: rgba(239, 68, 68, 0.2);
        color: #EF4444;
    }
    
    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #1F2937;
        border-radius: 4px 4px 0 0;
        padding: 0.5rem 1rem;
        color: white;
    }
    .stTabs [aria-selected="true"] {
        background-color: #2563EB;
    }
    
    /* File uploader styling */
    .stFileUploader label {
        color: white !important;
    }
    .uploadedFile {
        background-color: #1F2937 !important; 
        color: white !important;
        border: 1px solid #374151 !important;
    }
    
    /* Footer copyright */
    .footer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background-color: #1F2937;
        color: rgba(255, 255, 255, 0.5);
        text-align: center;
        padding: 5px;
        font-size: 12px;
        z-index: 999;
    }
    
    /* Additional spacing and layout fixes */
    .block-container {
        max-width: 1200px;
        padding-top: 2rem;
        padding-bottom: 3rem;
    }
    
    /* Columns spacing */
    [data-testid="column"] {
        padding: 0 0.5rem;
    }
    
    /* Remove padding from stExpander headers */
    .st-expander [data-testid="stExpander"] {
        padding: 0;
    }
    
    /* Adjust margins for section headers */
    [data-testid="stVerticalBlock"] > div:first-child > div:first-child h3 {
        margin-top: 0 !important;
    }
</style>

<!-- Footer -->
<div class="footer">
    Property of Andre Øvstedal - Please do not distribute
</div>
""", unsafe_allow_html=True)

# App layout
def main():
    # Title and introduction
    st.title("Altera JHA Assistant")
    st.markdown("Generate Job Hazard Analysis documents for maritime operations")
    
    # Create columns for layout
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # Job Details section - using proper container
        with st.container():
            st.subheader("Job Details")
            
            # Vessel Name
            vessel_name = st.text_input("Vessel Name")
            
            # Location
            location = st.selectbox("Work Location", [
                "Engine Room", "Wheelhouse Top", "Pump Room", "Thruster Room", 
                "HPR Trunk", "Main Deck", "Bridge", "Accommodation", 
                "Forward Mast", "Tank Entry", "MOB", "FFLB", "Gangway", "Other"
            ])
            
            # Task Name
            task_name = st.text_input("Task Name")
            
            # Job Description
            job_desc = st.text_area("Job Description", height=100)
        
            # Work Types section - using proper container
            with st.container():
                st.subheader("Work Types")
    
                # Create a vertical list of work types
                st.markdown("### Select applicable work types:")
    
                # Initialize selected work types list
                selected_work_types = []
    
                # Working at Height category
                st.markdown("**Working at Height:**")
                working_aloft = st.checkbox("Working Aloft")
                overside = st.checkbox("Overside")
    
                # Critical Equipment category
                st.markdown("**Critical Equipment:**")
                critical_equipment = st.checkbox("Critical Equipment")
                critical_system = st.checkbox("Critical System")
                
                # Other Work Types category
                st.markdown("**Other Work Types:**")
                enclosed_space = st.checkbox("Enclosed Space Entry")
                hot_work = st.checkbox("Hot Work")
                cold_work = st.checkbox("Cold Work")
                loto = st.checkbox("LOTO")
                
                # Add selections to the list
                if working_aloft:
                    selected_work_types.append("Working Aloft")
                if overside:
                    selected_work_types.append("Overside")
                if critical_equipment:
                    selected_work_types.append("Critical Equipment")
                if critical_system:
                    selected_work_types.append("Critical System")
                if enclosed_space:
                    selected_work_types.append("Enclosed Space Entry")
                if hot_work:
                    selected_work_types.append("Hot Work")
                if cold_work:
                    selected_work_types.append("Cold Work")
                if loto:
                    selected_work_types.append("LOTO")
        
        # Generate Button
        if st.button("Generate JHA", type="primary"):
            if not job_desc:
                st.error("Please enter a job description")
            elif not selected_work_types:
                st.error("Please select at least one work type")
            else:
                with st.spinner("Generating JHA..."):
                    st.session_state.jha_data = create_jha(job_desc, selected_work_types, location)
                    add_message_to_chat("Generated JHA successfully. Please review and let me know if you need any changes.", "assistant")
        
        # Equipment Manuals section - using proper container
        with st.container():
            st.subheader("Equipment Manuals")
            
            # Upload button
            uploaded_file = st.file_uploader("Upload Equipment Manual", type=["pdf"])
            if uploaded_file is not None:
                # Save and process the uploaded file
                filename = uploaded_file.name
                if filename not in [manual['name'] for manual in st.session_state.equipment_manuals]:
                    # Create temporary file and save
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                        tmp_file.write(uploaded_file.getvalue())
                        filepath = tmp_file.name
                    
                    # Add to the list
                    st.session_state.equipment_manuals.append({
                        'name': filename,
                        'path': filepath,
                        'size': uploaded_file.size
                    })
                    
                    st.success(f"Uploaded {filename}")
            
            # Display list of uploaded manuals
            if st.session_state.equipment_manuals:
                st.markdown("**Uploaded Manuals:**")
                for manual in st.session_state.equipment_manuals:
                    st.markdown(f"- {manual['name']} ({round(manual['size']/1024, 1)} KB)")
            else:
                st.markdown("No manuals uploaded yet...")
        
        # Chat section - using proper container
        with st.container():
            st.subheader("Assistant Chat")
            
            # Display chat history
            for message in st.session_state.chat_history:
                if message['sender'] == 'user':
                    st.markdown(f"<div class='chat-message user-message'><strong>You:</strong> {message['text']}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='chat-message assistant-message'><strong>Assistant:</strong> {message['text']}</div>", unsafe_allow_html=True)
            
            # Chat input
            chat_input = st.text_input("Type your message...")
            if st.button("Send"):
                if chat_input:
                    # Add user message to chat
                    add_message_to_chat(chat_input, "user")
                    
                    if st.session_state.jha_data:
                        # Process the message
                        with st.spinner("Processing..."):
                            try:
                                # Update JHA based on message
                                updated_jha = update_jha_with_message(st.session_state.jha_data, chat_input)
                                
                                # Update session state
                                st.session_state.jha_data = updated_jha
                                
                                # Add response to chat
                                add_message_to_chat("I've updated the JHA based on your request. Please review the changes.", "assistant")
                                
                                # Rerun to update display
                                st.experimental_rerun()
                            except Exception as e:
                                add_message_to_chat(f"I couldn't update the JHA: {str(e)}", "assistant")
                    else:
                        add_message_to_chat("Please generate a JHA first before sending requests.", "assistant")
                        
                    # Clear input (this won't actually clear it in Streamlit, but helps with the logic)
                    chat_input = ""
    
    with col2:
        # JHA Display section - using proper container
        with st.container():
            st.subheader("Generated JHA")
            
            # Display JHA if available
            if st.session_state.jha_data:
                # Create tabs for different views
                tab1, tab2 = st.tabs(["Formatted View", "JSON View"])
                
                with tab1:
                    # Display formatted JHA
                    display_formatted_jha(st.session_state.jha_data, vessel_name, task_name)
                    
                    # Download button
                    if vessel_name and task_name:
                        doc_bytes = create_jha_document(st.session_state.jha_data, vessel_name, task_name)
                        if doc_bytes:
                            st.download_button(
                                label="Download JHA Document",
                                data=doc_bytes,
                                file_name=f"JHA_{task_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                
                with tab2:
                    # Display JSON view
                    st.json(st.session_state.jha_data)
            else:
                st.markdown("No JHA generated yet. Fill in the details and click 'Generate JHA'.")

def add_message_to_chat(text, sender):
    """Add a message to the chat history"""
    st.session_state.chat_history.append({
        'text': text,
        'sender': sender
    })

def display_formatted_jha(jha_data, vessel_name, task_name):
    """Display JHA in a formatted view"""
    
    # Header info
    st.markdown(f"### JHA for: {task_name}")
    st.markdown(f"**Vessel:** {vessel_name} | **Date:** {datetime.now().strftime('%d.%m.%Y')}")
    st.markdown("---")
    
    # Display steps
    st.markdown("### Steps and Hazards")
    
    for i, step in enumerate(jha_data.get('steps', []), 1):
        with st.expander(f"Step {i}: {step.get('description', '')[:50]}...", expanded=True if i <= 3 else False):
            st.markdown(f"**Description:**")
            st.markdown(step.get('description', ''))
            
            st.markdown("**Hazards:**")
            hazards = step.get('hazards', {})
            if isinstance(hazards, dict):
                st.markdown(f"i. {hazards.get('potential_hazards', '')}")
                st.markdown(f"ii. {hazards.get('who_affected', '')}")
                st.markdown(f"iii. {hazards.get('how_occurs', '')}")
            else:
                st.markdown(str(hazards))
            
            st.markdown("**Controls:**")
            controls = step.get('controls', [])
            if isinstance(controls, list):
                for control in controls:
                    st.markdown(f"- {control}")
            else:
                st.markdown(str(controls))
            
            # Risk level
            risk = step.get('risk_level', {})
            if isinstance(risk, dict):
                c = risk.get('consequence', 1)
                l = risk.get('likelihood', 1)
                r = c * l
                st.markdown(f"**Risk Level:** {get_risk_badge(c, l)}", unsafe_allow_html=True)
    
    # Permits
    if jha_data.get('permits_required'):
        st.markdown("---")
        st.markdown("### Required Permits")
        for permit in jha_data['permits_required']:
            st.markdown(f"- {permit}")
    
    # Special considerations
    if jha_data.get('special_considerations'):
        st.markdown("---")
        st.markdown("### Special Considerations")
        considerations = jha_data['special_considerations']
        if isinstance(considerations, dict):
            for key, value in considerations.items():
                st.markdown(f"**{key.replace('_', ' ').title()}:** {value}")
        elif isinstance(considerations, list):
            for consideration in considerations:
                st.markdown(f"- {consideration}")
        else:
            st.markdown(str(considerations))

if __name__ == "__main__":
    main()
